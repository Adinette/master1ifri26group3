"""Generateur du Rapport de Projet - Phase 12.

Phase 12 : refonte UI/UX du portail (dark mode utilisateur, dashboard analytics,
audit de coherence frontend, restauration du menu microservices).

Convention : pas d accents pour rester compatible avec tous les terminaux et
les versions Word anciennes. Les chaines longues sont decoupees pour la lisibilite.
"""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / 'docs' / 'Rapport_Projet_TWM_Phase12.docx'

PHASE12_BRANCH = 'phase-12-ui-theme-dashboard'
PHASE12_BASE = 'phase-11-inventory-fix'


PHASE12_KPIS = [
    ['Fichiers modifies (UI/theme)', '17'],
    ['Nouveaux fichiers (theme + toggle)', '2'],
    ['Pages dashboard auditees', '12/12'],
    ['Bugs critiques corriges', '1 (sidebar overlap)'],
    ['Modes de theme disponibles', '3 (light / dark / system)'],
    ['Graphes ajoutes au dashboard', '2 (histogramme + camembert)'],
    ['Dependance graphique ajoutee', '0 (SVG natif)'],
]


PHASE12_DELIVERABLES = [
    ['app/lib/theme.tsx', 'Nouveau',
     'ThemeProvider + hook useTheme avec persistance localStorage et suivi prefers-color-scheme.'],
    ['app/components/ThemeToggle.tsx', 'Nouveau',
     'Composant dropdown accessible (3 options) reutilisable navbar publique et dashboard.'],
    ['app/globals.css', 'Mise a jour',
     '@custom-variant dark Tailwind v4 + variables CSS dark pour bascule par classe.'],
    ['app/layout.tsx', 'Mise a jour',
     'Script inline anti-FOUC dans <head> + suppressHydrationWarning sur <html>.'],
    ['app/providers.tsx', 'Mise a jour',
     'Wrapper ThemeProvider autour de SessionProvider.'],
    ['app/dashboard/page.tsx', 'Refonte',
     'Vrai dashboard analytics : KPIs, histogramme, camembert, recents, etat infra.'],
    ['app/dashboard/settings/page.tsx', 'Refonte',
     'Selecteur de theme interactif + carte etat microservices live + tech info.'],
    ['app/components/DashboardSidebar.tsx', 'Mise a jour',
     'Section Systeme avec entree Microservices + corrections dark mode.'],
    ['app/components/DashboardNavbar.tsx', 'Mise a jour',
     'Integration ThemeToggle + dark mode aligne sur le reste.'],
    ['app/components/Navbar.tsx', 'Mise a jour',
     'Suppression complete des inline styles, conversion Tailwind + dark mode.'],
    ['app/dashboard/layout.tsx', 'Fix critique',
     'Ajout lg:ml-64 pour eviter le chevauchement contenu/sidebar fixe en desktop.'],
    ['app/dashboard/profile/page.tsx', 'Mise a jour',
     'Badge En attente avec variantes dark.'],
    ['app/dashboard/services/page.tsx', 'Mise a jour',
     'Palette unifiee zinc + dark mode.'],
    ['app/dashboard/reporting/page.tsx', 'Mise a jour',
     'Badge filtre dark mode aligne.'],
    ['app/front/auth/login/page.tsx', 'Mise a jour',
     'Suppression styles inline, ajout variantes dark sur les fonds et inputs.'],
]


AUDIT_ROWS = [
    ['Sidebar fixe vs contenu', 'Contenu sous la sidebar 256px sur desktop',
     'Ajout lg:ml-64 sur le wrapper principal du layout dashboard'],
    ['Inline styles Navbar publique', 'Hardcoded white, aucun support dark',
     'Conversion 100% Tailwind avec variantes dark'],
    ['Logo mobile sidebar', 'Texte SFMC en text-lg dans h-8 w-8 (overflow)',
     'Reduction a text-[8px] pour tenir dans la pastille'],
    ['Brand sidebar dark', 'dark:text-zinc-400 quasi invisible',
     'Passage a dark:text-white'],
    ['Palette incoherente', 'Mix slate-500 / zinc-500 / gray-*',
     'Standardisation sur zinc-* avec variantes dark'],
    ['Cartes accueil dashboard', 'bg-orange-50 etc. sans variantes dark',
     'Ajout dark:bg-*-900/10 sur toutes les cartes quick action'],
    ['Badge profile En attente', 'bg-yellow-100 sans dark',
     'Ajout dark:bg-yellow-900/30 + dark:text-yellow-400'],
]


THEME_FEATURES = [
    'Trois modes selectionnables : Clair, Sombre, Systeme.',
    'Persistance du choix utilisateur via localStorage (cle theme).',
    'Suivi reactif de prefers-color-scheme quand le mode Systeme est actif.',
    'Script anti-FOUC inline dans <head> qui applique .dark avant l hydratation React.',
    'suppressHydrationWarning sur <html> pour eviter l avertissement de mismatch.',
    'Dropdown accessible : aria-haspopup, aria-expanded, role menuitemradio.',
    'Fermeture par clic exterieur et touche Escape.',
    'Toggle present dans la navbar publique (desktop + mobile) et la navbar dashboard.',
]


DASHBOARD_FEATURES = [
    'Hero compact avec salutation utilisateur et CTA actions rapides.',
    'Cinq cartes KPI : Commandes, Stock, Factures, Production, Revenus FCFA.',
    'Histogramme SVG natif des commandes par statut (en attente / validees / expediees / annulees).',
    'Camembert SVG natif des factures (payees vs en attente) avec legende et pourcentages.',
    'Quatre listes recents : commandes, factures, notifications, stock critique.',
    'Widget infrastructure : pastilles live de Kong, RabbitMQ et microservices.',
    'Loader, gestion d erreur, et badge donnees partielles si un service repond 503.',
    'Aucune dependance ajoutee : tous les graphes sont en SVG natif.',
]


SETTINGS_SECTIONS = [
    ['Apparence', 'Trois cartes cliquables (Clair / Sombre / Systeme).',
     'Indique le mode applique ; carte active mise en evidence.'],
    ['General', 'Langue, devise (FCFA / EUR).',
     'Placeholder coherent prepare pour internationalisation.'],
    ['Etat du systeme', 'Statut live Kong + RabbitMQ + microservices.',
     'Bouton Actualiser et lien vers la page de detail.'],
    ['Informations techniques', 'Versions des stacks principales.',
     'Next.js, React, Prisma, NextAuth, PostgreSQL, Tailwind.'],
]


SIDEBAR_CHANGES = [
    'Nouvelle section Systeme dans la sidebar.',
    'Entree Microservices avec icone satellite vers /dashboard/services.',
    'Acces aux statuts d infrastructure depuis n importe quelle page du dashboard.',
    'Correctifs visuels : palette zinc, dark mode, logo mobile redimensionne.',
]


VALIDATION_ROWS = [
    ['Compilation Next.js', 'Lancement npm run dev', 'Ready en 1.2 s sans erreur'],
    ['Page dashboard', 'GET http://localhost:3000/dashboard', '200 OK (compile + hydratation)'],
    ['Page settings', 'GET http://localhost:3000/dashboard/settings', '200 OK avec toggle interactif'],
    ['Page login', 'GET http://localhost:3000/front/auth/login', '200 OK en clair et sombre'],
    ['Theme persistance', 'localStorage.theme entre rafraichissements', 'Choix conserve'],
    ['Anti-FOUC', 'Rechargement en mode sombre', 'Aucun flash blanc visible'],
    ['Reactivite Systeme', 'Bascule OS clair vers sombre', 'Interface bascule sans rechargement'],
]


PROBLEM_ROWS = [
    ['Chevauchement sidebar/contenu sur desktop',
     'Sidebar fixed sans marge gauche sur le wrapper',
     'Ajout lg:ml-64 sur le conteneur flex principal'],
    ['Flash blanc au chargement en mode sombre',
     'Classe .dark appliquee apres l hydratation React',
     'Script inline synchrone dans <head> avant tout rendu'],
    ['Avertissement React hydration mismatch',
     'Mutation du <html> par le script avant rendu',
     'Ajout suppressHydrationWarning sur la balise html'],
    ['Lien client_fetch_error sur le toggle',
     'useTheme appele hors du provider',
     'Wrapper ThemeProvider dans app/providers.tsx'],
    ['Tailwind v4 et dark mode par classe',
     'Par defaut Tailwind v4 utilise prefers-color-scheme',
     '@custom-variant dark (&:where(.dark, .dark *)) dans globals.css'],
]


LOCAL_UPDATE_STEPS = [
    'Recuperer la branche : git fetch origin puis git checkout phase-12-ui-theme-dashboard.',
    'Installer les dependances : npm install.',
    'Demarrer l infrastructure : npm run dev:infra.',
    'Demarrer le frontend et les microservices : npm run dev:full.',
    'Ouvrir http://localhost:3000/dashboard et tester le toggle dans la navbar.',
]


CLONE_STEPS = [
    'Cloner le depot puis se placer dans tp_twm.',
    'Checkout la branche phase-12-ui-theme-dashboard.',
    'Installer les dependances racine et services.',
    'Generer les Prisma clients.',
    'Demarrer l infra puis l ensemble des services.',
    'Verifier les pages cles : /, /catalogue, /dashboard, /dashboard/settings.',
]


# ----------------------------------------------------------------------
# Iteration finale Phase 12 (avril 2026) - consolidation metier + perf
# ----------------------------------------------------------------------


FINAL_KPIS = [
    ['Endpoints saga annulation ajoutes', '4'],
    ['Cache backend ajoutes', '3 (reporting 30s, services-status 10s, kong-status 10s)'],
    ['Timeouts microservices', '800 ms (reporting) / 600 ms (probes)'],
    ['Composants memoises (React.memo)', '4 (KpiCard, BarChart, DonutChart, InfraDot)'],
    ['Helper client de cache', '1 (app/lib/client-cache.ts)'],
    ['Pages adoptant le cache client', '3 (orders, billing, production)'],
    ['Incoherences de marque corrigees', '5 (sidebar, layout, login)'],
    ['Erreurs lint introduites', '0'],
]


FINAL_DELIVERABLES = [
    ['services/inventory/lib/stock-workflow.ts', 'Mise a jour',
     'Ajout releaseStockForOrder + publication stock.updated sur tout mouvement (in/out/adjust/release).'],
    ['services/inventory/app/api/stock/release/route.ts', 'Nouveau',
     'Endpoint POST pour liberer le stock reserve a une commande.'],
    ['services/billing/lib/billing-store.ts', 'Mise a jour',
     'Ajout cancelInvoice et cancelInvoicesByOrder pour annulation des factures non payees.'],
    ['services/billing/app/api/invoices/cancel-by-order/[orderId]/route.ts', 'Nouveau',
     'Endpoint POST pour annuler toutes les factures d une commande.'],
    ['services/order/app/api/orders/[id]/cancel/route.ts', 'Nouveau',
     'Saga d annulation : libere stock, annule factures, met a jour statut, publie order.cancelled.'],
    ['services/order/app/api/orders/[id]/route.ts', 'Mise a jour',
     'Acceptation des statuts cancelled et failed sur PUT.'],
    ['services/reporting/app/api/dashboard/route.ts', 'Mise a jour',
     'Timeout 800 ms par appel + exposition KPIs cancelledOrders et deliveredOrders.'],
    ['app/api/orders/[id]/cancel/route.ts', 'Nouveau',
     'Proxy d annulation, transmet le cookie de session vers Order Service.'],
    ['app/api/reporting/dashboard/route.ts', 'Mise a jour',
     'Cache memoire 30 s avec stale-while-error en cas de panne reporting.'],
    ['app/api/services-status/route.ts', 'Mise a jour',
     'Cache 10 s, timeout 600 ms par probe, coalescing in-flight des requetes.'],
    ['app/api/kong-status/route.ts', 'Mise a jour',
     'Cache 10 s, timeout 600 ms, coalescing in-flight.'],
    ['app/lib/client-cache.ts', 'Nouveau',
     'Helper cachedJson + invalidate : cache memoire 5 s et deduplication des fetches GET.'],
    ['app/dashboard/page.tsx', 'Mise a jour',
     'React.memo sur KpiCard/BarChart/DonutChart/InfraDot, fetch infra differe de 200 ms, fix mutation DonutChart.'],
    ['app/dashboard/orders/page.tsx', 'Mise a jour',
     'Bouton Annuler + KPI Annulees + adoption client-cache + invalidation cross-page.'],
    ['app/dashboard/billing/page.tsx', 'Mise a jour',
     'Adoption client-cache, invalidation des caches orders/invoices/payments apres mutation.'],
    ['app/dashboard/production/page.tsx', 'Mise a jour',
     'Adoption client-cache, invalidation /api/stock au passage en completed.'],
    ['app/components/DashboardSidebar.tsx', 'Fix coherence',
     'Bloc desktop : Benin ERP -> SFMC Benin (alignement avec le reste du portail).'],
    ['app/layout.tsx', 'Fix coherence',
     'Metadata : SFMC Benin avec accent + Plateforme d operations corrige.'],
    ['app/front/auth/login/page.tsx', 'Fix coherence',
     'Titre Connexion - SFMC Benin -> Connexion -- SFMC Benin (avec accent).'],
]


CANCEL_SAGA_STEPS = [
    'Frontend ouvre POST /api/orders/:id/cancel via le bouton Annuler de /dashboard/orders.',
    'Le proxy Next.js relaie la requete vers Order Service avec le cookie de session.',
    'Order Service appelle Inventory : POST /api/stock/release (libere le stock reserve).',
    'Order Service appelle Billing : POST /api/invoices/cancel-by-order/:orderId (annule les factures non payees).',
    'Order Service met le statut de la commande a cancelled et publie order.cancelled sur RabbitMQ.',
    'Notification consumer ecoute order.cancelled et notifie le client.',
    'Reporting voit le KPI cancelledOrders augmenter et stock.updated mettre a jour les niveaux.',
]


PERF_OPTIMIZATIONS = [
    ['Proxy reporting', 'Cache memoire 30 s + stale-while-error',
     'Le dashboard supporte une panne courte du service reporting sans afficher d erreur.'],
    ['Appels micros depuis reporting', 'Timeout 800 ms (AbortSignal.timeout)',
     'Si un micro est lent, on bascule en mode degrade au lieu de bloquer la page.'],
    ['Probes infra (Kong + services)', 'Cache 10 s + coalescing in-flight + timeout 600 ms',
     '9 probes simultanees sont fusionnees en une seule requete reseau cache 10 s.'],
    ['Composants graphiques', 'React.memo (4 composants)',
     'Plus de re-render a chaque tick du SessionTimeoutManager (1 tick / s).'],
    ['Fetch infra dashboard', 'Differe de 200 ms apres le 1er paint',
     'Le reporting recoit la priorite reseau pour le 1er rendu de la page.'],
    ['Helper client', 'cachedJson + invalidate (5 s)',
     'Navigation orders / billing / production sans re-fetch si donnees deja en memoire.'],
    ['DonutChart', 'reduce au lieu de let offset (mutation)',
     'Conforme a la regle react-hooks/immutability de React 19 / Next 16.'],
]


BRANDING_FIXES = [
    ['app/components/DashboardSidebar.tsx (desktop)', 'Benin ERP', 'SFMC Benin'],
    ['app/layout.tsx (title)', 'SFMC Benin - Plateforme d operations', 'SFMC Benin -- Plateforme d operations (accents)'],
    ['app/layout.tsx (description)', '... pour SFMC Benin.', '... pour SFMC Benin. (avec accent)'],
    ['app/front/auth/login/page.tsx (h1 #1)', 'Connexion - SFMC Benin', 'Connexion -- SFMC Benin (accent)'],
    ['app/front/auth/login/page.tsx (h1 #2)', 'Connexion - SFMC Benin', 'Connexion -- SFMC Benin (accent)'],
]


FINAL_VALIDATION_ROWS = [
    ['Compilation TypeScript root', 'npx tsc --noEmit -p tsconfig.json', 'Exit 0'],
    ['Compilation TypeScript order/inventory/billing', 'npx tsc --noEmit -p services/<svc>/tsconfig.json', 'Exit 0'],
    ['Compilation TypeScript reporting', 'npx tsc --noEmit -p services/reporting/tsconfig.json', 'Exit 0 apres prisma generate'],
    ['Lint UI', 'npx eslint app/**/*.{ts,tsx}', '0 erreur introduite par cette livraison'],
    ['Saga annulation', 'Bouton Annuler sur /dashboard/orders puis V des KPIs',
     'Stock libere, facture annulee, statut cancelled, KPI cancelledOrders +1'],
    ['Cache reporting (resilience)', 'Couper reporting puis recharger /dashboard',
     'Affichage de l ancienne snapshot avec badge donnees partielles'],
    ['Coherence marque', 'Sidebar desktop + login + onglet navigateur',
     'SFMC Benin partout, plus aucune occurrence de Benin ERP'],
]


doc = Document()


def setup_styles() -> None:
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        heading = doc.styles[f'Heading {level}']
        heading.font.name = 'Calibri'
        heading.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)


def add_title_page() -> None:
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Rapport de Projet - Phase 12')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run(
        'Theme utilisateur, dashboard analytics et audit de coherence UI'
    )
    sub.font.size = Pt(15)
    sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Projet : ').bold = True
    meta.add_run('tp_twm | SFMC Benin | Groupe 5\n')
    meta.add_run('Depot GitHub : ').bold = True
    meta.add_run('github.com/Adinette/tp_twm\n')
    meta.add_run('Branche active : ').bold = True
    meta.add_run(f'{PHASE12_BRANCH}\n')
    meta.add_run('Base technique : ').bold = True
    meta.add_run(f'{PHASE12_BASE}\n')
    meta.add_run('Date : ').bold = True
    meta.add_run(datetime.now().strftime('%d/%m/%Y'))

    doc.add_page_break()


def add_table(headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            table.rows[row_index].cells[col_index].text = value


def add_bullet(text: str) -> None:
    doc.add_paragraph(text, style='List Bullet')


def add_code(text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def build_document() -> None:
    doc.add_heading('Table des matieres', level=1)
    toc_items = [
        '1. Objectif de la phase 12',
        '2. Ce qui change par rapport a la phase 11',
        '3. Ce qui a ete livre',
        '4. Fichiers livres',
        '5. Audit de coherence UI',
        '6. Validation effectuee',
        '7. Problemes rencontres et solutions',
        '8. Parametrage et mode operatoire',
        '9. Iteration finale Phase 12 (avril 2026)',
        '10. Suite recommandee',
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # 1. Objectif
    doc.add_heading('1. Objectif de la phase 12', level=1)
    doc.add_paragraph(
        'La phase 12 apporte une refonte UI ciblee du portail SFMC Benin. '
        'Trois axes : permettre a l utilisateur de choisir son theme '
        '(clair, sombre, systeme), transformer le tableau de bord en page '
        'd analyse avec graphes et indicateurs, et restaurer dans le menu '
        'l acces direct au monitoring des microservices.'
    )
    add_bullet('Donner le controle du theme a l utilisateur final, sans rechargement.')
    add_bullet('Rendre le dashboard exploitable pour la prise de decision.')
    add_bullet('Standardiser la coherence visuelle (palette, dark mode, responsivite).')
    add_bullet('Remettre le suivi infrastructure dans le menu de navigation.')

    # 2. Comparaison
    doc.add_paragraph()
    doc.add_heading('2. Ce qui change par rapport a la phase 11', level=1)
    add_table(
        ['Sujet', 'Phase 11', 'Phase 12'],
        [
            ['Perimetre principal',
             'Fiabilite Inventory et flux event-driven',
             'Refonte UI : theme, dashboard analytics, audit'],
            ['Theme',
             'Suivi automatique du systeme via media query',
             'Choix utilisateur explicite + persistance localStorage'],
            ['Tableau de bord',
             'Page d accueil avec liens d acces rapide',
             'Stats consolidees, histogramme, camembert, recents'],
            ['Menu lateral',
             'Pas d acces direct aux microservices',
             'Section Systeme avec entree Microservices'],
            ['Graphes',
             'Aucun',
             'Histogramme + camembert en SVG natif'],
            ['Dependances ajoutees',
             'Aucune',
             'Aucune (pas de chart.js, pas de recharts)'],
        ],
    )

    doc.add_paragraph('Indicateurs quantifies de la phase 12 :')
    add_table(['Indicateur', 'Valeur'], PHASE12_KPIS)

    # 3. Livraison
    doc.add_paragraph()
    doc.add_heading('3. Ce qui a ete livre', level=1)

    doc.add_heading('3.1 Theme utilisateur (light / dark / system)', level=2)
    for item in THEME_FEATURES:
        add_bullet(item)
    doc.add_paragraph('Bascule Tailwind v4 vers le mode classe :')
    add_code(
        '/* app/globals.css */\n'
        '@custom-variant dark (&:where(.dark, .dark *));\n\n'
        '.dark {\n'
        '  --background: #09090B;       /* zinc-950 */\n'
        '  --foreground: #FAFAFA;       /* zinc-50 */\n'
        '  --surface: #18181B;          /* zinc-900 */\n'
        '  --surface-muted: #27272A;    /* zinc-800 */\n'
        '}'
    )
    doc.add_paragraph('Script anti-FOUC inline dans <head> :')
    add_code(
        '(function(){\n'
        "  try {\n"
        "    var s = localStorage.getItem('theme');\n"
        "    var t = (s==='dark'||s==='light'||s==='system') ? s : 'system';\n"
        "    var d = t==='dark' ||\n"
        "      (t==='system' && window.matchMedia('(prefers-color-scheme: dark)').matches);\n"
        "    if (d) document.documentElement.classList.add('dark');\n"
        '  } catch(e) {}\n'
        '})();'
    )

    doc.add_heading('3.2 Tableau de bord analytics', level=2)
    for item in DASHBOARD_FEATURES:
        add_bullet(item)
    doc.add_paragraph('Sources de donnees consommees :')
    add_code(
        'GET /api/reporting/dashboard   -> KPIs et listes recentes\n'
        'GET /api/kong-status            -> statut Kong + RabbitMQ\n'
        'GET /api/services-status        -> statut individuel des microservices'
    )

    doc.add_heading('3.3 Page Parametres refondue', level=2)
    add_table(
        ['Section', 'Contenu', 'Detail'],
        SETTINGS_SECTIONS,
    )

    doc.add_heading('3.4 Sidebar et menu Microservices', level=2)
    for item in SIDEBAR_CHANGES:
        add_bullet(item)

    # 4. Fichiers
    doc.add_paragraph()
    doc.add_heading('4. Fichiers livres', level=1)
    add_table(['Fichier', 'Type', 'Apport'], PHASE12_DELIVERABLES)

    # 5. Audit
    doc.add_paragraph()
    doc.add_heading('5. Audit de coherence UI', level=1)
    doc.add_paragraph(
        'Une passe systematique a ete realisee sur l ensemble des pages du portail '
        'pour detecter les incoherences de palette, les manques de variantes dark '
        'et les bugs responsives. Les principaux ecarts trouves et corriges :'
    )
    add_table(['Anomalie', 'Diagnostic', 'Correction'], AUDIT_ROWS)

    # 6. Validation
    doc.add_paragraph()
    doc.add_heading('6. Validation effectuee', level=1)
    add_table(['Controle', 'Execution', 'Resultat'], VALIDATION_ROWS)

    doc.add_paragraph('Commandes representatives utilisees pendant la verification :')
    add_code(
        'npm run dev\n'
        'Invoke-WebRequest -Uri http://localhost:3000/dashboard -UseBasicParsing\n'
        'Invoke-WebRequest -Uri http://localhost:3000/dashboard/settings -UseBasicParsing'
    )

    # 7. Problemes
    doc.add_paragraph()
    doc.add_heading('7. Problemes rencontres et solutions', level=1)
    add_table(['Incident', 'Cause', 'Resolution'], PROBLEM_ROWS)

    # 8. Mode operatoire
    doc.add_paragraph()
    doc.add_heading('8. Parametrage et mode operatoire', level=1)

    doc.add_heading('8.1 Si le projet est deja present en local', level=2)
    for item in LOCAL_UPDATE_STEPS:
        add_bullet(item)
    doc.add_paragraph('Commandes de base :')
    add_code(
        'git fetch origin\n'
        'git checkout phase-12-ui-theme-dashboard\n'
        'npm install\n'
        'npm run dev:infra\n'
        'npm run dev:full'
    )

    doc.add_heading('8.2 Si le projet est un nouveau clone', level=2)
    for item in CLONE_STEPS:
        add_bullet(item)
    doc.add_paragraph('Commandes de base :')
    add_code(
        'git clone https://github.com/Adinette/tp_twm.git\n'
        'cd tp_twm\n'
        'git checkout phase-12-ui-theme-dashboard\n'
        'npm install\n'
        'npm run dev:infra\n'
        'npm run dev:full'
    )

    doc.add_heading('8.3 Verifications apres demarrage', level=2)
    add_bullet('Ouvrir http://localhost:3000 et basculer le theme depuis la navbar.')
    add_bullet('Ouvrir /dashboard et verifier l affichage des graphes (histogramme et camembert).')
    add_bullet('Ouvrir /dashboard/settings et tester les trois modes (Clair, Sombre, Systeme).')
    add_bullet('Ouvrir /dashboard/services et confirmer la disponibilite Kong, RabbitMQ et microservices.')
    add_bullet('Verifier la persistance du choix : recharger la page, le theme reste applique.')

    # 9bis. Iteration finale Phase 12
    doc.add_paragraph()
    doc.add_heading(
        '9. Iteration finale Phase 12 (avril 2026)', level=1
    )
    doc.add_paragraph(
        'Apres la livraison initiale de la Phase 12 (theme + dashboard + audit UI), '
        'une derniere passe a ete realisee sur trois axes critiques : la conformite '
        'metier (saga d annulation de commande), la performance perceptive du '
        'dashboard, et la coherence visuelle de la marque SFMC Benin a travers '
        'l ensemble du portail.'
    )

    doc.add_heading('9.1 Indicateurs de l iteration finale', level=2)
    add_table(['Indicateur', 'Valeur'], FINAL_KPIS)

    doc.add_heading('9.2 Saga d annulation de commande', level=2)
    doc.add_paragraph(
        'L annulation d une commande implique trois services dans l ordre. Chacun '
        'expose un endpoint dedie ; l Order Service joue le role d orchestrateur :'
    )
    for step in CANCEL_SAGA_STEPS:
        add_bullet(step)
    doc.add_paragraph('Endpoints introduits :')
    add_code(
        'POST /api/orders/:id/cancel              (Order, orchestrateur)\n'
        'POST /api/stock/release                  (Inventory)\n'
        'POST /api/invoices/cancel-by-order/:id   (Billing)\n'
        'POST /api/orders/:id/cancel              (Frontend proxy Next.js)'
    )
    doc.add_paragraph(
        'Le KPI cancelledOrders est expose par le reporting et affiche dans la '
        'carte Annulees du dashboard et de la page Commandes.'
    )

    doc.add_heading('9.3 Optimisations de performance', level=2)
    add_table(['Cible', 'Mecanisme', 'Effet'], PERF_OPTIMIZATIONS)
    doc.add_paragraph('Helper client introduit :')
    add_code(
        '// app/lib/client-cache.ts\n'
        'await cachedJson<Order[]>(\'/api/orders\')         // GET cache 5 s\n'
        'invalidate(\'/api/orders\')                         // apres POST/PUT/DELETE\n'
        'invalidate(\'/api/stock\') ; invalidate(\'/api/invoices\')  // saga annulation'
    )

    doc.add_heading('9.4 Coherence de la marque', level=2)
    doc.add_paragraph(
        'Un audit visuel a detecte une incoherence : la sidebar desktop du dashboard '
        'affichait "Benin ERP" tandis que tout le reste du portail utilise '
        '"SFMC Benin". Quatre points de divergence ont ete identifies et alignes :'
    )
    add_table(['Fichier', 'Avant', 'Apres'], BRANDING_FIXES)
    doc.add_paragraph(
        'Les identifiants techniques (sfmc-benin.bj, animation CSS sfmc-ticker, '
        'URL WhatsApp encodees) n ont volontairement pas ete modifies.'
    )

    doc.add_heading('9.5 Validation effectuee', level=2)
    add_table(['Controle', 'Execution', 'Resultat'], FINAL_VALIDATION_ROWS)

    doc.add_heading('9.6 Fichiers livres dans cette iteration', level=2)
    add_table(['Fichier', 'Type', 'Apport'], FINAL_DELIVERABLES)

    # 10. Suite
    doc.add_paragraph()
    doc.add_heading('10. Suite recommandee', level=1)
    add_bullet('Internationalisation des libelles (FR/EN) sur le selecteur de theme et le dashboard.')
    add_bullet('Ajout de filtres temporels sur le dashboard (jour / semaine / mois).')
    add_bullet('Sauvegarde du theme cote serveur dans le profil utilisateur (table users).')
    add_bullet('Tests Playwright pour la persistance du theme et le rendu des graphes.')
    add_bullet('Extension des graphes : courbe d evolution des commandes, repartition stock par entrepot.')
    add_bullet('Mode haute densite pour le dashboard (compact / detaille).')

    # Conclusion
    doc.add_paragraph()
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'La phase 12 livre une experience utilisateur plus riche et plus controlee : '
        'le portail respecte le choix de theme de chacun, le tableau de bord donne '
        'enfin une vue exploitable des donnees consolidees, et la coherence visuelle '
        'globale a ete restauree. Aucune dependance n a ete ajoutee : tous les graphes '
        'et le moteur de theme sont implementes en code natif React + SVG, ce qui '
        'preserve la legerete du bundle et la maintenabilite du projet.'
    )


def main() -> None:
    setup_styles()
    add_title_page()
    build_document()
    doc.save(OUTPUT)
    print(f'Rapport genere : {OUTPUT}')


if __name__ == '__main__':
    main()
