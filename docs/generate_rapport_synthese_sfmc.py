from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / 'docs' / 'Rapport_Synthese_SFMC_Benin.docx'

PROJECT_NAME = 'SFMC Benin'
REPOSITORY_NAME = 'tp_twm'
REPOSITORY_URL = 'github.com/Adinette/tp_twm'
CURRENT_BRANCH = 'main'

PHASE_ROWS = [
    ['1', 'fix/prisma-auth-setup', 'Stabilisation Prisma 7, NextAuth et premier dashboard', 'Integre'],
    ['2', 'feature/phase2-register-cleanup', 'Inscription utilisateur et nettoyage du schema Prisma', 'Integre'],
    ['3', 'feature/phase3-components-ui', 'Navbar, footer, sidebar et base UI du dashboard', 'Integre'],
    ['4', 'feature/phase4-google-oauth', 'Google OAuth2 et documentation associee', 'Integre'],
    ['5', 'feature/phase5-kong-gateway', 'Kong Gateway, RabbitMQ et socle services', 'Integre'],
    ['6', 'feature/phase6-microservices-core', 'Microservices coeur et premiers flux inter-services', 'Integre'],
    ['7', 'feature/phase7-microservices-advanced', 'Pipeline etendu avec reporting, billing, production et notification', 'Integre'],
    ['8', 'feature/phase8-front-reporting-deployment-prep', 'Consolidation frontend, supervision reelle, reporting dashboard', 'Integre'],
    ['9', 'feature/phase9-collaboration-release-hub', 'UI/UX renforcee et hub de versions pour les collaborateurs', 'Integre'],
    ['10', 'feature/phase10-idle-session-security', 'Expiration par inactivite, avertissement et prolongation de session', 'Integre sur main'],
]

SERVICE_ROWS = [
    ['Frontend principal', '3000', 'Portail principal, dashboard, BFF et auth web'],
    ['Auth Service', '3001', 'Gestion auth et logique associee'],
    ['User Service', '3002', 'Utilisateurs, roles et profils'],
    ['Product Service', '3003', 'Catalogue produit'],
    ['Inventory Service', '3004', 'Stocks et mouvements'],
    ['Order Service', '3005', 'Commandes et orchestration'],
    ['Production Service', '3006', 'Suivi production'],
    ['Billing Service', '3007', 'Facturation'],
    ['Notification Service', '3008', 'Notifications et evenements'],
    ['Reporting Service', '3009', 'Synthese dashboard et indicateurs'],
    ['RabbitMQ', '5672 / 15672', 'Broker AMQP et interface de supervision'],
    ['Kong Gateway', '8000 / 8001', 'API gateway et administration'],
]

FEATURE_ROWS = [
    ['Authentification', 'Credentials + Google OAuth2 + sessions JWT'],
    ['Securite session', 'Expiration absolue + idle timeout + modal d avertissement'],
    ['Dashboard', 'Accueil dashboard, navigation laterale, navbar contextuelle'],
    ['Utilisateurs', 'CRUD utilisateur present avec roles visibles'],
    ['Catalogue', 'Pages produits, stock, mouvements et reporting'],
    ['Observabilite', 'Etat des services, Kong et RabbitMQ'],
    ['Collaboration', 'Hub de versions et documentation de recuperation'],
]

CHECK_ROWS = [
    ['1', 'git checkout main && git pull origin main', 'Le depot local est synchronise avec la branche principale'],
    ['2', 'npm install', 'Les dependances racine sont installees'],
    ['3', 'npm run dev:infra', 'Kong, sa base et RabbitMQ sont demarres'],
    ['4', 'npm run dev:services', 'Les microservices demarrent ensemble avec Webpack'],
    ['5', 'npm run dev', 'Le frontend principal est disponible sur localhost:3000'],
    ['6', 'Ouvrir /dashboard et /dashboard/services', 'Le tableau de bord et les statuts sont visibles'],
    ['7', 'Tester login, produits, stock, reporting', 'Les parcours principaux repondent'],
]

RISK_ROWS = [
    ['Nom technique historique', 'Le dossier Git et le depot conservent encore le nom tp_twm', 'Utiliser SFMC Benin dans les documents et supports, tout en gardant tp_twm comme identifiant technique'],
    ['Conflits de ports', 'Un service deja demarre bloque son port au relancement', 'Arreter les serveurs existants avant un lancement groupe'],
    ['Connexion user a verifier', 'Un retour equipe signale un probleme de connexion avec certains utilisateurs crees', 'Verifier l alignement entre creation utilisateur, table users et logique de login'],
    ['RBAC incomplet', 'Les roles existent mais le controle d acces serveur reste partiel', 'Prioriser une phase 11 orientee administration et autorisations'],
]

NEXT_ROWS = [
    ['1', 'Personnaliser l accueil avec le contexte SFMC Benin et un contenu plus realiste'],
    ['2', 'Verifier et corriger le parcours complet de connexion des utilisateurs crees depuis la gestion users'],
    ['3', 'Finaliser un controle d acces RBAC cote session, BFF et middleware'],
    ['4', 'Nettoyer progressivement les anciens intitulés TWM dans la documentation visible'],
    ['5', 'Preparer la demo finale et un support de presentation court'],
]

doc = Document()


def setup_styles() -> None:
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        heading = doc.styles[f'Heading {level}']
        heading.font.name = 'Calibri'
        heading.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)


def add_title_page() -> None:
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Document de synthese - SFMC Benin')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run('Vision globale du projet, etat actuel de main et mode operatoire equipe')
    sub.font.size = Pt(15)
    sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Projet : ').bold = True
    meta.add_run(f'{PROJECT_NAME}\n')
    meta.add_run('Nom technique du depot : ').bold = True
    meta.add_run(f'{REPOSITORY_NAME}\n')
    meta.add_run('Depot GitHub : ').bold = True
    meta.add_run(f'{REPOSITORY_URL}\n')
    meta.add_run('Branche de reference : ').bold = True
    meta.add_run(f'{CURRENT_BRANCH}\n')
    meta.add_run('Date : ').bold = True
    meta.add_run(datetime.now().strftime('%d/%m/%Y'))

    doc.add_page_break()


def add_table(headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            table.rows[row_index].cells[col_index].text = value


def add_bullet(text: str) -> None:
    doc.add_paragraph(text, style='List Bullet')


def add_code(text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def build_document() -> None:
    doc.add_heading('Table des matieres', level=1)
    toc_items = [
        '1. Resume executif',
        '2. Identite du projet',
        '3. Etat actuel sur la branche main',
        '4. Synthese des phases 1 a 10',
        '5. Architecture et services',
        '6. Fonctionnalites visibles',
        '7. Recuperer et demarrer le projet',
        '8. Tests rapides de verification',
        '9. Points de vigilance',
        '10. Prochaines actions recommandees',
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    doc.add_heading('1. Resume executif', level=1)
    doc.add_paragraph(
        'SFMC Benin est une application web modulaire basee sur Next.js, Prisma, PostgreSQL, Kong et RabbitMQ. '
        'Le projet a evolue par phases successives jusqu a disposer d un frontend principal, de neuf services applicatifs, '
        'd un dashboard de supervision, d une authentification web et d une securite de session renforcee sur la branche main.'
    )
    add_bullet('Le nom de communication du projet doit etre SFMC Benin.')
    add_bullet('Le nom tp_twm reste aujourd hui l identifiant technique du depot et du dossier Git.')
    add_bullet('La base la plus recente a utiliser par l equipe est la branche main.')

    doc.add_paragraph()
    doc.add_heading('2. Identite du projet', level=1)
    add_table(
        ['Element', 'Valeur'],
        [
            ['Nom projet pour les rapports et la soutenance', PROJECT_NAME],
            ['Nom technique du depot', REPOSITORY_NAME],
            ['Depot GitHub', REPOSITORY_URL],
            ['Branche actuelle de reference', CURRENT_BRANCH],
            ['Style de livraison recommande', 'Travail par branche de phase puis reintegration sur main'],
        ],
    )

    doc.add_paragraph()
    doc.add_heading('3. Etat actuel sur la branche main', level=1)
    doc.add_paragraph(
        'La branche main concentre aujourd hui la base a presenter. Elle contient l application principale, les microservices, '
        'les proxies API du frontend, la surveillance des services et la securite de session basee sur l inactivite.'
    )
    add_bullet('Le dashboard principal existe avec une navigation plus structuree et des pages metier visibles.')
    add_bullet('Les services produits, stock, reporting, notification, billing, production et orders sont relies au socle principal.')
    add_bullet('Un hub de versions et des documents de phase existent pour capitaliser les livraisons precedentes.')
    add_bullet('La session utilisateur peut maintenant expirer apres inactivite avec avertissement avant deconnexion.')

    doc.add_paragraph()
    doc.add_heading('4. Synthese des phases 1 a 10', level=1)
    add_table(['Phase', 'Branche', 'Apport principal', 'Etat'], PHASE_ROWS)

    doc.add_paragraph()
    doc.add_heading('5. Architecture et services', level=1)
    doc.add_paragraph(
        'Le systeme repose sur un frontend principal qui joue aussi le role de BFF pour plusieurs routes API, '
        'puis sur un ensemble de microservices Next.js specialises et relies a PostgreSQL, RabbitMQ et Kong.'
    )
    add_table(['Composant', 'Port', 'Responsabilite'], SERVICE_ROWS)

    doc.add_paragraph()
    doc.add_heading('6. Fonctionnalites visibles', level=1)
    add_table(['Bloc', 'Contenu'], FEATURE_ROWS)

    doc.add_paragraph()
    doc.add_heading('7. Recuperer et demarrer le projet', level=1)
    doc.add_paragraph('Commandes conseillees pour une machine de developpement Windows :')
    add_code(
        'git fetch origin\n'
        'git checkout main\n'
        'git pull origin main\n\n'
        'npm install\n'
        'npm run dev:infra\n'
        'npm run dev:services\n\n'
        '# Dans un autre terminal pour le frontend principal\n'
        'npm run dev'
    )
    doc.add_paragraph('Si vous partez d une machine totalement libre et voulez tout lancer depuis la racine :')
    add_code('npm run dev:all')
    doc.add_paragraph(
        'Les scripts dev:services et dev:all ont ete prepares pour Windows en forçant Webpack, '
        'ce qui evite les erreurs Turbopack observees sur plusieurs services.'
    )

    doc.add_paragraph()
    doc.add_heading('8. Tests rapides de verification', level=1)
    add_table(['#', 'Verification', 'Resultat attendu'], CHECK_ROWS)

    doc.add_paragraph()
    doc.add_heading('9. Points de vigilance', level=1)
    add_table(['Sujet', 'Constat', 'Action recommandee'], RISK_ROWS)

    doc.add_paragraph()
    doc.add_heading('10. Prochaines actions recommandees', level=1)
    add_table(['Priorite', 'Action'], NEXT_ROWS)

    doc.add_paragraph()
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'Ce document sert de resume global pour l equipe. Il permet de presenter le projet sous le nom SFMC Benin, '
        'de rappeler que la branche de reference est main, de montrer les phases deja livrees et de donner un mode operatoire simple '
        'pour relancer le systeme. Il constitue aussi une base utile pour la soutenance et pour la preparation de la prochaine phase de securisation.'
    )


def main() -> None:
    setup_styles()
    add_title_page()
    build_document()
    doc.save(OUTPUT)
    print(f'Rapport genere : {OUTPUT}')


if __name__ == '__main__':
    main()