from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / 'docs' / 'Rapport_Projet_TWM_Phase8.docx'

doc = Document()


def setup_styles() -> None:
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        heading = doc.styles[f'Heading {level}']
        heading.font.name = 'Calibri'
        heading.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)


def add_title_page() -> None:
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Rapport de Projet — Phase 8')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run('Consolidation Frontend Phase 7 + Supervision réelle + Préparation au déploiement')
    sub.font.size = Pt(15)
    sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Projet : ').bold = True
    meta.add_run('tp_twm | SFMC Bénin | Groupe 5\n')
    meta.add_run('Dépôt GitHub : ').bold = True
    meta.add_run('github.com/Adinette/tp_twm\n')
    meta.add_run('Branche active : ').bold = True
    meta.add_run('feature/phase8-front-reporting-deployment-prep\n')
    meta.add_run('Base technique : ').bold = True
    meta.add_run('feature/phase7-microservices-advanced de Silas\n')
    meta.add_run('Date : ').bold = True
    meta.add_run(datetime.now().strftime('%d/%m/%Y'))

    doc.add_page_break()


def add_table(headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            table.rows[row_index].cells[col_index].text = value


def add_bullet(text: str) -> None:
    doc.add_paragraph(text, style='List Bullet')


def add_code(text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def build_document() -> None:
    doc.add_heading('Table des matières', level=1)
    toc_items = [
        '1. Où en sommes-nous ? — Bilan des 8 phases',
        '2. Stack technique complète',
        '3. Mettre à jour le projet sur sa machine',
        '4. Ce qui a été fait en Phase 8',
        '5. Tests à effectuer après récupération',
        '6. Problèmes rencontrés et solutions',
        '7. Proposition concrète pour la Phase 9',
        '8. Répartition des tâches — Équipe de 3 membres',
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    doc.add_heading('1. Où en sommes-nous ? — Bilan des 8 phases', level=1)
    doc.add_paragraph(
        'La base fonctionnelle du projet a été livrée en Phase 7 avec les 9 microservices et le pipeline RabbitMQ complet. '
        'La Phase 8 présentée ici prend appui sur cette base pour consolider la couche frontend principale, rendre la supervision réelle '
        'et préparer proprement le passage au déploiement.'
    )
    add_table(
        ['Phase', 'Branche', 'Ce qui a été fait', 'Statut'],
        [
            ['Phase 1', 'fix/prisma-auth-setup', 'Prisma 7 + auth JWT + dashboard initial', '✅ Mergé'],
            ['Phase 2', 'feature/phase2-register-cleanup', 'Inscription + nettoyage schéma Prisma', '✅ PR ouverte'],
            ['Phase 3', 'feature/phase3-components-ui', 'Navbar, footer, sidebar, dashboard', '✅ PR ouverte'],
            ['Phase 4', 'feature/phase4-google-oauth', 'Google OAuth2 + callback signIn + docs', '✅ PR ouverte'],
            ['Phase 5', 'feature/phase5-kong-gateway', 'Kong + RabbitMQ + page services', '✅ Disponible sur main'],
            ['Phase 6', 'feature/phase6-microservices-core', 'Services core + pipeline OrderCreated', '✅ Livrée'],
            ['Phase 7', 'feature/phase7-microservices-advanced', 'Production, Billing, Notification, Reporting + pipeline complet', '✅ Base de référence'],
            ['Phase 8', 'feature/phase8-front-reporting-deployment-prep', 'Consolidation frontend, supervision réelle, reporting dashboard, préparation GitHub/déploiement', '✅ En cours de livraison'],
        ],
    )

    doc.add_paragraph()
    doc.add_heading('2. Stack technique complète', level=1)
    add_table(
        ['Technologie', 'Version', 'Rôle'],
        [
            ['Next.js', '16.2.2', 'Frontend principal et microservices indépendants'],
            ['React', '19.2.4', 'UI côté client'],
            ['TypeScript', '5', 'Typage statique'],
            ['Prisma', '7.7.0', 'ORM sur toutes les bases PostgreSQL'],
            ['PostgreSQL', '18', 'Base de données principale et bases par service'],
            ['NextAuth.js', '4.24.13', 'Auth JWT + Google OAuth2'],
            ['Kong Gateway', '3.6', 'API gateway'],
            ['RabbitMQ', '3-management', 'Message broker'],
            ['Docker Compose', 'latest', 'Orchestration locale'],
        ],
    )

    doc.add_paragraph()
    doc.add_heading('3. Mettre à jour le projet sur sa machine', level=1)
    doc.add_paragraph('Pour récupérer la Phase 8, partir de la nouvelle branche et relancer les services utiles au frontend consolidé :')
    add_code(
        'git fetch origin\n'
        'git checkout feature/phase8-front-reporting-deployment-prep\n'
        'git pull origin feature/phase8-front-reporting-deployment-prep\n\n'
        'npm install\n'
        'docker-compose up -d\n\n'
        'cd services/product && npm run dev\n'
        'cd ..\\inventory && npm run dev\n'
        'cd ..\\reporting && npm run dev\n'
        'cd ..\\.. && npm run dev'
    )
    doc.add_paragraph('Si l’on veut valider le reporting complet, il faut aussi démarrer les services order, billing, production et notification, puis leurs routes /api/init lorsque nécessaire.')

    doc.add_paragraph()
    doc.add_heading('4. Ce qui a été fait en Phase 8', level=1)
    doc.add_heading('4.1 Consolidation frontend du dashboard principal', level=2)
    add_bullet('Ajout des pages dashboard Produits et Stock pour exposer Product Service et Inventory Service dans l’application principale.')
    add_bullet('Ajout d’une page dashboard Reporting branchée sur le service reporting via un proxy dédié.')
    add_bullet('Mise à jour de la sidebar et de la page dashboard pour rendre ces accès visibles et cohérents.')

    doc.add_heading('4.2 Proxies API dans le frontend principal', level=2)
    add_bullet('GET/POST /api/products → Product Service (3003).')
    add_bullet('GET/PUT/DELETE /api/products/[id] → Product Service (3003).')
    add_bullet('GET/POST /api/stock → Inventory Service (3004).')
    add_bullet('GET/POST /api/mouvements → Inventory Service (3004).')
    add_bullet('GET /api/reporting/dashboard → Reporting Service (3009).')

    doc.add_heading('4.3 Supervision réelle des microservices', level=2)
    add_bullet('Création de /api/services-status pour sonder réellement les 9 services au lieu de simuler leur état côté client.')
    add_bullet('Correction de /dashboard/services : les statuts affichés reflètent désormais la disponibilité réelle des endpoints des services.')
    add_bullet('Conservation de /api/kong-status pour l’état de Kong et RabbitMQ.')

    doc.add_heading('4.4 Préparation au déploiement', level=2)
    add_bullet('Création d’une branche dédiée Phase 8 pour isoler la consolidation frontend et les livrables documentaires.')
    add_bullet('Structuration du travail pour une future PR propre, séparée des documents anciens centrés Phase 6.')
    add_bullet('Production du rapport Phase 8 en s’appuyant sur le format documentaire utilisé par Silas en Phase 6 et 7.')

    doc.add_paragraph()
    doc.add_heading('5. Tests à effectuer après récupération', level=1)
    add_table(
        ['#', 'Action', 'Résultat attendu'],
        [
            ['1', 'git branch --show-current', 'feature/phase8-front-reporting-deployment-prep'],
            ['2', 'GET http://localhost:3003/api/products', 'Catalogue produits disponible'],
            ['3', 'GET http://localhost:3004/api/stock', 'Stocks disponibles'],
            ['4', 'GET http://localhost:3009/api/dashboard', 'Agrégation reporting disponible'],
            ['5', 'Ouvrir /dashboard/products', 'Liste produits chargée via /api/products'],
            ['6', 'Ouvrir /dashboard/stock', 'Stocks et mouvements visibles'],
            ['7', 'Ouvrir /dashboard/reporting', 'Résumé global chargé'],
            ['8', 'Ouvrir /dashboard/services', 'Statuts réels des 9 services affichés'],
            ['9', 'Créer un produit puis un mouvement', 'Chaîne front → proxy → microservice valide'],
        ],
    )
    doc.add_paragraph('Exemples de commandes PowerShell utiles :')
    add_code(
        'Invoke-RestMethod http://localhost:3003/api/products\n'
        'Invoke-RestMethod http://localhost:3004/api/stock\n'
        'Invoke-RestMethod http://localhost:3004/api/mouvements\n'
        'Invoke-RestMethod http://localhost:3009/api/dashboard\n'
        'Invoke-RestMethod http://localhost:8001/status'
    )

    doc.add_paragraph()
    doc.add_heading('6. Problèmes rencontrés et solutions', level=1)
    add_table(
        ['Problème', 'Cause', 'Solution'],
        [
            ['Supervision toujours hors ligne', 'La page services simulait offline localement', 'Création de /api/services-status avec sondes réelles'],
            ['Reporting invisible depuis le frontend principal', 'Le service reporting existait mais aucun proxy/page n’était exposé', 'Ajout de /api/reporting/dashboard et de /dashboard/reporting'],
            ['Routes dynamiques Next.js 16', 'params est asynchrone', 'Conserver le pattern await params dans les routes dynamiques'],
            ['Risque de bruit documentaire', 'Des fichiers locaux étaient encore centrés Phase 6', 'Isoler la livraison Phase 8 sur une branche dédiée et un nouveau rapport'],
        ],
    )

    doc.add_paragraph()
    doc.add_heading('7. Proposition concrète pour la Phase 9', level=1)
    doc.add_paragraph(
        'La Phase 9 doit se concentrer sur la mise en production et la finalisation du projet après cette consolidation frontend.'
    )
    add_table(
        ['Étape', 'Tâche', 'Responsable'],
        [
            ['1', 'Déployer le frontend principal sur Vercel', 'Membre 1'],
            ['2', 'Créer / configurer les bases cloud sur Railway ou Neon', 'Membre 1'],
            ['3', 'Finaliser les Dockerfiles et l’orchestration des 9 services', 'Membre 2'],
            ['4', 'Configurer les variables d’environnement et les URIs Google OAuth de production', 'Membre 1'],
            ['5', 'Lancer les tests end-to-end complets en environnement cible', 'Membres 1, 2 et 3'],
            ['6', 'Préparer les slides et la démonstration finale', 'Membre 3'],
        ],
    )

    doc.add_paragraph()
    doc.add_heading('8. Répartition des tâches — Équipe de 3 membres', level=1)
    add_table(
        ['Membre', 'Rôle', 'Apport Phase 8', 'Suite Phase 9'],
        [
            ['Membre 1', 'Lead / Infra', 'Cadrage de la branche Phase 8 et préparation déploiement', 'Vercel + Kong + variables prod'],
            ['Membre 2', 'Dev Fullstack', 'Consolidation API/proxy/front produits-stock-reporting', 'Packaging services + validation technique'],
            ['Membre 3', 'Frontend / QA', 'Vérification dashboard, supervision et rendu final', 'QA finale + slides + documentation'],
        ],
    )

    doc.add_paragraph()
    doc.add_heading('Conclusion et appel à l’action', level=1)
    doc.add_paragraph(
        'Cette Phase 8 ne remplace pas la Phase 7 : elle la consolide. La base microservices de Silas est conservée, '
        'mais le frontend principal devient enfin capable d’exposer Produits, Stock, Reporting et une supervision réelle des services. '
        'La suite logique est désormais le déploiement propre et la démonstration finale.'
    )
    add_bullet('Récupérer la branche Phase 8')
    add_bullet('Démarrer les services nécessaires')
    add_bullet('Vérifier les écrans dashboard/products, dashboard/stock, dashboard/reporting et dashboard/services')
    add_bullet('Ouvrir la PR et préparer le passage à la Phase 9')


def main() -> None:
    setup_styles()
    add_title_page()
    build_document()
    doc.save(OUTPUT)
    print(f'Rapport généré : {OUTPUT}')


if __name__ == '__main__':
    main()